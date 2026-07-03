"""
Vikturi AI — Streamlit web interface for the multi-agent ecosystem.
Run: streamlit run app.py
"""
from __future__ import annotations
import io
import json
import os
import re
import subprocess
import sys
import tempfile
import uuid
from datetime import datetime
from pathlib import Path

# Ensure project root is on sys.path when launched via `streamlit run`
sys.path.insert(0, str(Path(__file__).parent))

os.environ.setdefault("CHROMA_SERVER_NOFILE", "65536")
os.environ.setdefault("ANONYMIZED_TELEMETRY", "false")
os.environ.setdefault("CREWAI_STORAGE_DIR", "/tmp/vikturi_storage")

import streamlit as st


def _response_to_docx(text: str) -> bytes:
    """Convert AI response (basic markdown) to a .docx file."""
    from docx import Document
    from docx.shared import Pt

    # Strip agent-label prefix like "🎨 **DrewAI**\n\n"
    clean = re.sub(r'^[^\n]*\*\*[^*]+\*\*[^\n]*\n\n', '', text, count=1).strip()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(90)

    for line in clean.split('\n'):
        s = line.strip()
        if not s:
            doc.add_paragraph()
        elif s.startswith('### '):
            doc.add_heading(s[4:], level=3)
        elif s.startswith('## '):
            doc.add_heading(s[3:], level=2)
        elif s.startswith('# '):
            doc.add_heading(s[2:], level=1)
        elif s.startswith(('- ', '* ')):
            doc.add_paragraph(s[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', s):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', s), style='List Number')
        else:
            p = doc.add_paragraph()
            for part in re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', s):
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert .docx bytes to PDF bytes using LibreOffice headless."""
    work_dir = Path(tempfile.mkdtemp())
    src = work_dir / "output.docx"
    src.write_bytes(docx_bytes)
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(work_dir), str(src)],
        check=True, capture_output=True, timeout=120,
    )
    return (work_dir / "output.pdf").read_bytes()


def _convert_document(file_bytes: bytes, suffix: str, original_name: str) -> tuple[bytes, str, str]:
    """Convert DOCX→PDF or PDF→DOCX. Returns (output_bytes, output_filename, mime_type)."""
    import subprocess
    import time

    work_dir = Path(tempfile.mkdtemp())
    stem = Path(original_name).stem

    if suffix == ".docx":
        # DOCX → PDF via LibreOffice headless
        src = work_dir / f"{stem}.docx"
        src.write_bytes(file_bytes)
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(work_dir), str(src)],
            check=True, capture_output=True, timeout=120,
        )
        out = work_dir / f"{stem}.pdf"
        return out.read_bytes(), f"{stem}.pdf", "application/pdf"

    else:
        # PDF → DOCX via pdf2docx
        from pdf2docx import Converter
        src = work_dir / f"{stem}.pdf"
        out = work_dir / f"{stem}.docx"
        src.write_bytes(file_bytes)
        cv = Converter(str(src))
        cv.convert(str(out))
        cv.close()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return out.read_bytes(), f"{stem}.docx", mime


def _render_message(content: str) -> None:
    """Renders a message, handling image/file markers:
    - 🖼️ IMAGE:<path>  → local file via st.image() (HF/Pollinations downloaded locally)
    - POLLINATIONS_IMG:<url> → fetched server-side and rendered as bytes
    - 📄 FILE:<path> → generated/improved code file, offered as a download button
    """
    # ── Generated code file (Dimelis) ─────────────────────────────────
    if "📄 FILE:" in content:
        import mimetypes

        before, rest = content.split("📄 FILE:", 1)
        if before.strip():
            st.markdown(before)
        first_line, *remaining = rest.split("\n", 1)
        file_path = Path(first_line.strip())
        if file_path.exists():
            display_name = file_path.name.split("_", 1)[-1]
            mime, _ = mimetypes.guess_type(display_name)
            st.download_button(
                f"⬇️ Descargar {display_name}",
                data=file_path.read_bytes(),
                file_name=display_name,
                mime=mime or "application/octet-stream",
                key=f"dl_{file_path.name}",
            )
        else:
            st.warning(f"Archivo no encontrado: `{file_path}`")
        if remaining and remaining[0].strip():
            st.markdown(remaining[0])
        return

    # ── HF / downloaded local image ──────────────────────────────────
    if "🖼️ IMAGE:" in content:
        before, rest = content.split("🖼️ IMAGE:", 1)
        if before.strip():
            st.markdown(before)
        first_line, *remaining = rest.split("\n", 1)
        img_path = Path(first_line.strip())
        if img_path.exists():
            st.image(str(img_path), use_container_width=True)
        else:
            st.warning(f"Imagen no encontrada: `{img_path}`")
        if remaining and remaining[0].strip():
            st.markdown(remaining[0])
        return

    # ── Pollinations image ────────────────────────────────────────────
    if "POLLINATIONS_IMG:" in content:
        before, rest = content.split("POLLINATIONS_IMG:", 1)
        if before.strip():
            st.markdown(before)
        first_line, *remaining = rest.split("\n", 1)
        url = first_line.strip()
        try:
            import requests as _req
            with st.spinner("🎨 Generando imagen…"):
                resp = _req.get(url, timeout=90)
            if resp.status_code == 200:
                st.image(resp.content, use_container_width=True)
            else:
                st.warning(f"Pollinations devolvió status {resp.status_code}. Intenta de nuevo.")
        except Exception as e:
            st.warning(f"No se pudo cargar la imagen: {e}")
        if remaining and remaining[0].strip():
            st.markdown(remaining[0])
        return

    # ── Plain text / markdown ─────────────────────────────────────────
    st.markdown(content)


# ── Conversation persistence helpers ─────────────────────────────────
_CONV_DIR = Path("/tmp/vikturi_conversations")


def _conv_path(conv_id: str) -> Path:
    _CONV_DIR.mkdir(parents=True, exist_ok=True)
    return _CONV_DIR / f"{conv_id}.json"


def _save_conv(conv_id: str, messages: list, title: str = "") -> None:
    path = _conv_path(conv_id)
    existing = json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}
    data = {
        "id": conv_id,
        "title": title or existing.get("title", "Nueva conversación"),
        "messages": messages,
        "created_at": existing.get("created_at", datetime.now().isoformat()),
        "updated_at": datetime.now().isoformat(),
    }
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _load_conv(conv_id: str) -> list:
    path = _conv_path(conv_id)
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8")).get("messages", [])


def _list_convs() -> list[dict]:
    _CONV_DIR.mkdir(parents=True, exist_ok=True)
    convs = []
    for p in sorted(_CONV_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            convs.append({
                "id": data["id"],
                "title": data.get("title", "Sin título"),
                "updated_at": data.get("updated_at", ""),
            })
        except Exception:
            pass
    return convs[:20]


def _auto_title(messages: list) -> str:
    for msg in messages:
        if msg.get("role") == "user":
            text = re.sub(r"^🎤 \[Voz transcrita\]: ", "", msg.get("content", "")).strip()
            if text:
                return text[:40] + ("…" if len(text) > 40 else "")
    return "Nueva conversación"


def _fmt_time(iso: str) -> str:
    try:
        dt = datetime.fromisoformat(iso)
        today = datetime.now().date()
        if dt.date() == today:
            return dt.strftime("Hoy %H:%M")
        return dt.strftime("%d %b %H:%M")
    except Exception:
        return ""


# ── Page config (must be first Streamlit call) ────────────────────────
st.set_page_config(
    page_title="Vikturi AI",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ────────────────────────────────────────────────────────
st.markdown("""
<style>
  .vk-title {
    font-size: 2.6rem;
    font-weight: 900;
    background: linear-gradient(135deg, #7C6EFA 0%, #3ECFCF 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    line-height: 1.1;
    margin: 0;
  }
  .vk-subtitle {
    font-size: 0.78rem;
    letter-spacing: 4px;
    text-transform: uppercase;
    color: #888;
    margin-top: 4px;
    margin-bottom: 20px;
  }
  .agent-pill {
    display: flex;
    align-items: center;
    gap: 10px;
    background: rgba(255,255,255,0.04);
    border-left: 3px solid #2ECC71;
    border-radius: 8px;
    padding: 9px 12px;
    margin-bottom: 7px;
  }
  .agent-pill .name { font-weight: 700; font-size: 0.9rem; }
  .agent-pill .desc { font-size: 0.75rem; color: #999; }
  hr { border-color: rgba(255,255,255,0.08) !important; }
  footer { visibility: hidden; }
  #MainMenu { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ── Auth gate ─────────────────────────────────────────────────────────
_APP_PASSWORD = os.getenv("APP_PASSWORD", "")
if _APP_PASSWORD and not st.session_state.get("authenticated"):
    st.markdown('<p class="vk-title">⚡ Vikturi AI</p>', unsafe_allow_html=True)
    st.markdown('<p class="vk-subtitle">The Multi-Agent Lore</p>', unsafe_allow_html=True)
    st.markdown("---")
    _pwd_input = st.text_input("🔑 Contraseña", type="password", placeholder="Ingresa la contraseña...")
    if st.button("Entrar", type="primary", use_container_width=True):
        if _pwd_input == _APP_PASSWORD:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("Contraseña incorrecta.")
    st.stop()

# ── Header ────────────────────────────────────────────────────────────
st.markdown('<p class="vk-title">⚡ Vikturi AI</p>', unsafe_allow_html=True)
st.markdown('<p class="vk-subtitle">The Multi-Agent Lore</p>', unsafe_allow_html=True)

# ── Sidebar ───────────────────────────────────────────────────────────
with st.sidebar:
    # ── Conversation history ──────────────────────────────────────────
    st.markdown("### 💬 Conversaciones")

    if st.button("➕ Nueva conversación", use_container_width=True, type="primary", key="new_conv_top"):
        if st.session_state.get("messages"):
            _save_conv(
                st.session_state["conv_id"],
                st.session_state.messages,
                _auto_title(st.session_state.messages),
            )
        st.session_state["conv_id"] = str(uuid.uuid4())[:8]
        st.session_state.messages = []
        for k in ("pending_image_bytes", "pending_image_suffix", "_pending_file_id",
                  "pending_voice_text", "_pending_audio_id",
                  "pending_doc_bytes", "pending_doc_suffix", "pending_doc_name", "_pending_doc_id",
                  "pending_video_bytes", "pending_video_suffix", "_pending_video_id",
                  "last_assistant_response"):
            st.session_state.pop(k, None)
        st.rerun()

    _convs = _list_convs()
    _current_id = st.session_state.get("conv_id", "")
    for _conv in _convs:
        _is_active = _conv["id"] == _current_id
        _label = f"{'▶ ' if _is_active else ''}{_conv['title']}"
        _caption = _fmt_time(_conv.get("updated_at", ""))
        _col_btn, _col_del = st.columns([5, 1])
        if _col_btn.button(
            _label, key=f"conv_{_conv['id']}",
            use_container_width=True,
            type="primary" if _is_active else "secondary",
            help=_caption,
        ):
            if not _is_active:
                if st.session_state.get("messages"):
                    _save_conv(_current_id, st.session_state.messages, _auto_title(st.session_state.messages))
                st.session_state["conv_id"] = _conv["id"]
                st.session_state.messages = _load_conv(_conv["id"])
                st.rerun()
        if _col_del.button("🗑", key=f"del_{_conv['id']}", help="Eliminar"):
            _conv_path(_conv["id"]).unlink(missing_ok=True)
            if _is_active:
                st.session_state["conv_id"] = str(uuid.uuid4())[:8]
                st.session_state.messages = []
            st.rerun()

    st.markdown("---")
    st.markdown("### 🟢 Estado del Sistema")
    st.markdown("---")

    _AGENTS = [
        ("🎨", "DrewAI",      "Visual, imágenes, TikTok"),
        ("🌵", "Gasp Tree",   "Marketing digital, campañas, banners"),
        ("💪", "Vigor AI",    "Nutrición deportiva, entrenamiento, biomecánica"),
        ("📚", "Yvannia AI",  "Tutorías paso a paso"),
        ("💻", "Dimelis AI",  "Código, estructura, PEP8"),
        ("🔍", "Teriania",    "Investigación web verificada"),
    ]
    for icon, name, desc in _AGENTS:
        st.markdown(
            f'<div class="agent-pill">'
            f'<span style="font-size:1.3rem">{icon}</span>'
            f'<div><div class="name">{name}</div><div class="desc">{desc}</div></div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    st.markdown("---")
    training_mode = st.toggle(
        "⚙️ Modo Entrenamiento",
        value=False,
        help="Lee los archivos de context/ para actualizar el briefing de los agentes",
    )

    st.markdown("---")

    # Session stats — fail silently if no history yet
    try:
        from memory.session_memory import get_stats
        stats = get_stats()
        st.markdown("### 📊 Estadísticas")
        col_a, col_b = st.columns(2)
        col_a.metric("Total", stats.get("total", 0))
        col_b.metric("Sesiones", stats.get("trainings", 0))
        if stats.get("last_date"):
            st.caption(f"Última actividad: {stats['last_date'][:10]}")
    except Exception:
        st.caption("Sin historial todavía.")

    # ── Download last response as document ───────────────────────────
    if "last_assistant_response" in st.session_state:
        st.markdown("---")
        st.markdown("### 💾 Descargar respuesta")
        st.caption("Convierte la última respuesta a documento.")
        _col_w, _col_p = st.columns(2)
        _last = st.session_state["last_assistant_response"]
        _docx_bytes = _response_to_docx(_last)
        _docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        _col_w.download_button(
            "⬇️ Word", data=_docx_bytes,
            file_name="vikturi_respuesta.docx", mime=_docx_mime,
            use_container_width=True,
        )
        try:
            _pdf_bytes = _docx_bytes_to_pdf(_docx_bytes)
            _col_p.download_button(
                "⬇️ PDF", data=_pdf_bytes,
                file_name="vikturi_respuesta.pdf", mime="application/pdf",
                use_container_width=True,
            )
        except Exception:
            _col_p.caption("PDF no disponible")

    # ── Image uploader for DrewAI ─────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📷 Imagen")
    st.caption("Sube una foto y DrewAI la analizará con tu próximo mensaje.")

    if "uploader_key" not in st.session_state:
        st.session_state["uploader_key"] = 0

    uploaded_file = st.file_uploader(
        "Subir imagen",
        type=["jpg", "jpeg", "png", "webp"],
        label_visibility="collapsed",
        key=f"img_upload_{st.session_state['uploader_key']}",
    )
    if uploaded_file is not None:
        file_id = f"{uploaded_file.name}_{uploaded_file.size}"
        if st.session_state.get("_pending_file_id") != file_id:
            st.session_state["pending_image_bytes"] = uploaded_file.getbuffer().tobytes()
            st.session_state["pending_image_suffix"] = Path(uploaded_file.name).suffix or ".jpg"
            st.session_state["_pending_file_id"] = file_id
        st.image(uploaded_file, use_container_width=True, caption="✅ Lista — envía tu mensaje")
    else:
        for k in ("pending_image_bytes", "pending_image_suffix", "_pending_file_id"):
            st.session_state.pop(k, None)

    # ── Voice input ───────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 🎤 Mensaje de voz")

    if "audio_uploader_key" not in st.session_state:
        st.session_state["audio_uploader_key"] = 0

    _hf_token = os.getenv("HF_TOKEN", "")
    if not _hf_token:
        st.warning("HF_TOKEN es requerido para transcribir voz.")

    _recorded_audio = st.audio_input(
        "Grabar mensaje",
        key=f"audio_rec_{st.session_state['audio_uploader_key']}",
    )
    _uploaded_audio = st.file_uploader(
        "O sube un archivo de audio",
        type=["mp3", "wav", "m4a", "ogg", "webm"],
        label_visibility="collapsed",
        key=f"audio_upload_{st.session_state['audio_uploader_key']}",
    )

    _audio_source = _recorded_audio or _uploaded_audio
    if _audio_source is not None and _hf_token:
        _audio_id = f"{getattr(_audio_source, 'name', 'recorded')}_{_audio_source.size}"
        if st.session_state.get("_pending_audio_id") != _audio_id:
            try:
                from huggingface_hub import InferenceClient as _HFClient
                _hf_client = _HFClient(token=_hf_token)
                _transcription = _hf_client.automatic_speech_recognition(
                    _audio_source.getvalue(),
                    model="openai/whisper-large-v3",
                )
                _transcribed_text = getattr(_transcription, "text", str(_transcription))
                st.session_state["pending_voice_text"] = f"🎤 [Voz transcrita]: {_transcribed_text}"
                st.session_state["_pending_audio_id"] = _audio_id
            except Exception as _exc:
                st.error(f"Error al transcribir: {_exc}")
    elif _audio_source is None:
        for k in ("pending_voice_text", "_pending_audio_id"):
            st.session_state.pop(k, None)

    if "pending_voice_text" in st.session_state:
        st.caption(f"✅ {st.session_state['pending_voice_text'][:60]}…")

    # ── Document uploader ─────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📄 Documento / Código")
    st.caption(
        "Sube un PDF, DOCX, TXT o un archivo de código (HTML, Python, JS, Java, "
        "CSS, etc.) para que Dimelis lo revise, mejore o lo uses como base."
    )

    if "doc_uploader_key" not in st.session_state:
        st.session_state["doc_uploader_key"] = 0

    _uploaded_doc = st.file_uploader(
        "Subir documento",
        type=[
            "pdf", "docx", "txt",
            "html", "htm", "md", "markdown",
            "py", "js", "jsx", "ts", "tsx",
            "css", "scss",
            "java", "kt",
            "c", "cpp", "h", "hpp",
            "go", "rb", "php", "rs",
            "sh", "bash",
            "json", "xml", "yaml", "yml",
            "sql",
        ],
        label_visibility="collapsed",
        key=f"doc_upload_{st.session_state['doc_uploader_key']}",
    )
    if _uploaded_doc is not None:
        _doc_id = f"{_uploaded_doc.name}_{_uploaded_doc.size}"
        if st.session_state.get("_pending_doc_id") != _doc_id:
            st.session_state["pending_doc_bytes"] = _uploaded_doc.getbuffer().tobytes()
            st.session_state["pending_doc_suffix"] = Path(_uploaded_doc.name).suffix.lower()
            st.session_state["pending_doc_name"] = _uploaded_doc.name
            st.session_state["_pending_doc_id"] = _doc_id
        st.caption(f"✅ {_uploaded_doc.name} — envía tu mensaje")

        from crew.ecosystem_simple import is_code_file_too_large
        if is_code_file_too_large(
            st.session_state["pending_doc_bytes"], st.session_state["pending_doc_suffix"]
        ):
            _size_kb = len(st.session_state["pending_doc_bytes"]) // 1024
            st.info(
                f"⚠️ Este archivo (~{_size_kb} KB) es grande para que Dimelis lo "
                "reescriba completo en una sola respuesta (el modelo tiene un límite "
                "de ~128,000 tokens de salida). Dimelis hará una revisión dirigida "
                "por secciones en vez de reescribir todo. Si quieres una parte "
                "específica reescrita, indícalo en tu mensaje."
            )
    else:
        for k in ("pending_doc_bytes", "pending_doc_suffix", "pending_doc_name", "_pending_doc_id"):
            st.session_state.pop(k, None)

    # ── Video uploader ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 🎬 Video")
    st.caption("DrewAI transcribe el audio y analiza fotogramas. Máx. ~100 MB.")

    if "video_uploader_key" not in st.session_state:
        st.session_state["video_uploader_key"] = 0

    _uploaded_video = st.file_uploader(
        "Subir video",
        type=["mp4", "mov", "avi", "webm", "mkv"],
        label_visibility="collapsed",
        key=f"video_upload_{st.session_state['video_uploader_key']}",
    )
    if _uploaded_video is not None:
        _vid_id = f"{_uploaded_video.name}_{_uploaded_video.size}"
        if st.session_state.get("_pending_video_id") != _vid_id:
            st.session_state["pending_video_bytes"] = _uploaded_video.getbuffer().tobytes()
            st.session_state["pending_video_suffix"] = Path(_uploaded_video.name).suffix.lower()
            st.session_state["_pending_video_id"] = _vid_id
        st.caption(f"✅ {_uploaded_video.name} — envía tu mensaje")
    else:
        for k in ("pending_video_bytes", "pending_video_suffix", "_pending_video_id"):
            st.session_state.pop(k, None)

    # ── Document converter ────────────────────────────────────────────
    st.markdown("---")
    with st.expander("🔄 Convertir documento"):
        st.caption("DOCX → PDF  ·  PDF → DOCX")
        _conv_file = st.file_uploader(
            "Sube el archivo a convertir",
            type=["pdf", "docx"],
            label_visibility="collapsed",
            key="conv_uploader",
        )
        if _conv_file is not None:
            _conv_suffix = Path(_conv_file.name).suffix.lower()
            _conv_label = "Convertir a PDF" if _conv_suffix == ".docx" else "Convertir a Word"
            if st.button(_conv_label, use_container_width=True, type="primary"):
                with st.spinner("Convirtiendo…"):
                    try:
                        _conv_result, _conv_out_name, _conv_mime = _convert_document(
                            _conv_file.getbuffer().tobytes(), _conv_suffix, _conv_file.name
                        )
                        st.download_button(
                            label=f"⬇️ Descargar {_conv_out_name}",
                            data=_conv_result,
                            file_name=_conv_out_name,
                            mime=_conv_mime,
                            use_container_width=True,
                        )
                    except Exception as _e:
                        st.error(f"Error al convertir: {_e}")

    # ── Clear conversation (bottom shortcut) ─────────────────────────
    st.markdown("---")
    if st.button("🗑️ Borrar conversación actual", use_container_width=True, type="secondary"):
        _conv_path(st.session_state.get("conv_id", "")).unlink(missing_ok=True)
        st.session_state["conv_id"] = str(uuid.uuid4())[:8]
        st.session_state.messages = []
        for k in ("pending_image_bytes", "pending_image_suffix", "_pending_file_id",
                  "pending_voice_text", "_pending_audio_id",
                  "pending_doc_bytes", "pending_doc_suffix", "pending_doc_name", "_pending_doc_id",
                  "pending_video_bytes", "pending_video_suffix", "_pending_video_id",
                  "last_assistant_response"):
            st.session_state.pop(k, None)
        st.session_state["uploader_key"] = st.session_state.get("uploader_key", 0) + 1
        st.session_state["audio_uploader_key"] = st.session_state.get("audio_uploader_key", 0) + 1
        st.session_state["doc_uploader_key"] = st.session_state.get("doc_uploader_key", 0) + 1
        st.session_state["video_uploader_key"] = st.session_state.get("video_uploader_key", 0) + 1
        st.rerun()

    st.markdown("---")
    st.caption("Vikturi AI · Multi-Agent Ecosystem · v1.1")

# ── Chat session state ────────────────────────────────────────────────
if "conv_id" not in st.session_state:
    st.session_state["conv_id"] = str(uuid.uuid4())[:8]

if "messages" not in st.session_state:
    st.session_state.messages = []

# Welcome message on first load
if not st.session_state.messages:
    st.session_state.messages.append({
        "role": "assistant",
        "content": (
            "¡Hola! Soy **Vikturi AI**, tu ecosistema multi-agente. "
            "Puedo delegarte a:\n\n"
            "- 💻 **Dimelis** para código y proyectos\n"
            "- 📚 **Yvannia** para aprender algo paso a paso\n"
            "- 🔍 **Teriania** para investigar con fuentes reales\n"
            "- 🎨 **DrewAI** para creatividad visual e imágenes\n\n"
            "También puedes **subir una imagen** desde el panel lateral para que DrewAI la analice.\n\n"
            "¿Con qué empezamos?"
        ),
    })

# ── Render chat history ───────────────────────────────────────────────
for msg in st.session_state.messages:
    avatar = "⚡" if msg["role"] == "assistant" else "🧑"
    with st.chat_message(msg["role"], avatar=avatar):
        _render_message(msg["content"])

# ── Chat input ────────────────────────────────────────────────────────
_placeholder = (
    "Modo entrenamiento activo — escribe qué procesar..."
    if training_mode
    else "Escribe tu pregunta para Vikturi AI..."
)

if prompt := st.chat_input(_placeholder):
    # Collect pending attachments
    pending_voice = st.session_state.get("pending_voice_text")
    pending_img: str | None = None
    if "pending_image_bytes" in st.session_state:
        suffix = st.session_state.get("pending_image_suffix", ".jpg")
        tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        tmp.write(st.session_state["pending_image_bytes"])
        tmp.close()
        pending_img = tmp.name
    pending_doc_bytes = st.session_state.get("pending_doc_bytes")
    pending_doc_suffix = st.session_state.get("pending_doc_suffix")
    pending_doc_name = st.session_state.get("pending_doc_name")
    pending_video_bytes = st.session_state.get("pending_video_bytes")
    pending_video_suffix = st.session_state.get("pending_video_suffix")
    pending_video_name = st.session_state.get("_pending_video_id", "")

    # Build the effective prompt (voice takes priority for text)
    effective_prompt = f"{pending_voice}\n\n{prompt}".strip() if pending_voice else prompt

    # 1. Show user bubble
    with st.chat_message("user", avatar="🧑"):
        if pending_voice:
            st.markdown(pending_voice)
        if pending_video_bytes:
            st.caption(f"🎬 Video adjunto — {len(pending_video_bytes) // 1024} KB")
        if pending_img:
            st.image(pending_img, use_container_width=True)
        if pending_doc_name:
            st.caption(f"📄 {pending_doc_name}")
        st.markdown(prompt)
    st.session_state.messages.append({"role": "user", "content": effective_prompt})

    # 2. Run the ecosystem — priority: video > image > document > voice/text
    with st.chat_message("assistant", avatar="⚡"):
        with st.status("🧠 Analizando y delegando al agente correcto…", expanded=False) as status_box:
            try:
                from crew.ecosystem_simple import run_simple

                chat_history = st.session_state.messages[:-1]

                if pending_video_bytes:
                    result = run_simple(
                        effective_prompt,
                        training_mode=training_mode,
                        chat_history=chat_history,
                        video_bytes=pending_video_bytes,
                        video_suffix=pending_video_suffix or ".mp4",
                    )
                elif pending_img:
                    result = run_simple(
                        effective_prompt,
                        training_mode=training_mode,
                        chat_history=chat_history,
                        image_path=pending_img,
                    )
                elif pending_doc_bytes:
                    result = run_simple(
                        effective_prompt,
                        training_mode=training_mode,
                        chat_history=chat_history,
                        doc_bytes=pending_doc_bytes,
                        doc_suffix=pending_doc_suffix,
                    )
                else:
                    result = run_simple(
                        effective_prompt,
                        training_mode=training_mode,
                        chat_history=chat_history,
                    )

                # Clear all consumed attachments and reset uploaders
                for k in ("pending_image_bytes", "pending_image_suffix", "_pending_file_id"):
                    st.session_state.pop(k, None)
                for k in ("pending_voice_text", "_pending_audio_id"):
                    st.session_state.pop(k, None)
                for k in ("pending_doc_bytes", "pending_doc_suffix", "pending_doc_name", "_pending_doc_id"):
                    st.session_state.pop(k, None)
                for k in ("pending_video_bytes", "pending_video_suffix", "_pending_video_id"):
                    st.session_state.pop(k, None)
                st.session_state["uploader_key"] = st.session_state.get("uploader_key", 0) + 1
                st.session_state["audio_uploader_key"] = st.session_state.get("audio_uploader_key", 0) + 1
                st.session_state["doc_uploader_key"] = st.session_state.get("doc_uploader_key", 0) + 1
                st.session_state["video_uploader_key"] = st.session_state.get("video_uploader_key", 0) + 1

                status_box.update(label="✅ Respuesta lista", state="complete", expanded=False)
            except Exception as exc:
                result = (
                    f"❌ **Error del ecosistema**\n\n"
                    f"```\n{exc}\n```\n\n"
                    f"Revisa que `ANTHROPIC_API_KEY` esté configurada en los secrets."
                )
                status_box.update(label="❌ Error", state="error", expanded=True)

        _render_message(result)

    st.session_state.messages.append({"role": "assistant", "content": result})
    st.session_state["last_assistant_response"] = result
    _save_conv(
        st.session_state["conv_id"],
        st.session_state.messages,
        _auto_title(st.session_state.messages),
    )
    st.rerun()
